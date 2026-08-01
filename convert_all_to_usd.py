#!/usr/bin/env python3
"""
AQUILA Sovereign Platform — Complete USD Converter
==================================================
Converts ALL files in the AQUILA Sovereign platform directory to a single
comprehensive OpenUSD stage:

  - Aquila_Master_BOM_Reconciled_v2 7-5-26.xlsx → BOM components as Mesh prims
  - *.docx engineering specs → EngineeringDesign prims with metadata
  - export_aquila_usd.py (existing) → reference for BOM geometry generation
  - aquila_sovereign_platform.usda (existing) → base stage to extend

Output: Aquila_Sovereign_Complete.usda

Usage:
    python convert_all_to_usd.py
    python convert_all_to_usd.py --output Aquila_Sovereign_Complete.usda
    python convert_all_to_usd.py --z-up   # Z-up for Isaac Sim
"""
from __future__ import annotations

import argparse
import math
import os
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

import docx
import openpyxl

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

BASE_DIR = Path(r"F:\OMNI\Vern\AQUILA Sovereign platform")
BOM_FILE = BASE_DIR / "Aquila_Master_BOM_Reconciled_v2 7-5-26 (version 1).xlsx"
OUTPUT_DEFAULT = BASE_DIR / "Aquila_Sovereign_Complete.usda"

# ---------------------------------------------------------------------------
# BOM Parsing
# ---------------------------------------------------------------------------

def parse_bom_xlsx(xlsx_path: Path) -> list[dict]:
    """Parse the Master BOM xlsx into component dictionaries."""
    wb = openpyxl.load_workbook(str(xlsx_path), data_only=True)
    ws = wb.active

    # Find header row (contains 'ID')
    header_row = None
    for r in range(1, 20):
        if str(ws.cell(r, 1).value or '').strip() == 'ID':
            header_row = r
            break
    if header_row is None:
        print("ERROR: Could not find BOM header row", file=sys.stderr)
        return []

    headers = [str(ws.cell(header_row, c).value or '').strip() for c in range(1, 13)]

    components = []
    for r in range(header_row + 1, ws.max_row + 1):
        bom_id = str(ws.cell(r, 1).value or '').strip()
        if not bom_id or bom_id.startswith('#') or bom_id.startswith('Phase'):
            continue

        comp = {}
        for c, h in enumerate(headers, 1):
            val = ws.cell(r, c).value
            if val is not None:
                comp[h] = str(val).strip()

        if comp.get('ID'):
            # Skip NOTE/summary rows
            bid = comp['ID']
            if bid.startswith('NOTE') or bid.startswith('note') or bid.startswith('TOTAL') or bid.startswith('SUMMARY'):
                continue
            # Parse dimensions
            dims_str = comp.get('Dimensions (mm)', '')
            dims = parse_dims(dims_str)
            comp['dims_mm'] = dims
            comp['weight_g'] = parse_number(comp.get('Weight (g)', '0'))
            comp['qty'] = int(parse_number(comp.get('Qty', '1')))
            comp['unit_cost'] = parse_number(comp.get('Unit Cost (USD)', '0'))
            comp['ext_cost'] = parse_number(comp.get('Ext. Cost (USD)', '0'))
            components.append(comp)

    return components


def parse_dims(s: str) -> tuple[float, float, float]:
    """Parse dimension strings like '215×90×75' or 'Ø300×270'."""
    s = s.replace('×', 'x').replace('Ø', '').replace('*', 'x')
    parts = re.findall(r'[\d.]+', s)
    if len(parts) >= 3:
        return (float(parts[0]), float(parts[1]), float(parts[2]))
    elif len(parts) == 2:
        return (float(parts[0]), float(parts[0]), float(parts[1]))  # cylinder
    elif len(parts) == 1:
        return (float(parts[0]), float(parts[0]), float(parts[0]))
    return (50.0, 50.0, 50.0)  # default


def parse_number(s: str) -> float:
    """Parse numbers with commas."""
    s = str(s).replace(',', '').replace('~', '').replace('$', '').strip()
    try:
        return float(s)
    except ValueError:
        return 0.0


# ---------------------------------------------------------------------------
# DOCX Parsing
# ---------------------------------------------------------------------------

def parse_docx_spec(docx_path: Path) -> dict:
    """Parse an engineering spec .docx into a structured dict."""
    try:
        doc = docx.Document(str(docx_path))
    except Exception as e:
        return {'filename': docx_path.name, 'error': str(e)}

    title = ''
    for p in doc.paragraphs:
        if p.text.strip():
            title = p.text.strip()
            break

    # Extract all text paragraphs
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Extract tables
    tables = []
    for t in doc.tables:
        table_data = []
        for row in t.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    # Try to find a Document ID
    doc_id = ''
    for table in tables:
        for row in table:
            if row and 'Document ID' in row[0]:
                doc_id = row[1] if len(row) > 1 else ''
                break

    # Extract TRIZ principle if present
    triz = ''
    for p in paragraphs:
        if 'TRIZ' in p:
            triz = p[:200]
            break

    # Extract date if present
    date = ''
    for p in paragraphs:
        if p.startswith('Date:'):
            date = p.replace('Date:', '').strip()[:30]
            break

    return {
        'filename': docx_path.name,
        'title': title,
        'doc_id': doc_id,
        'triz': triz,
        'date': date,
        'paragraphs': paragraphs,
        'tables': tables,
        'n_paragraphs': len(paragraphs),
        'n_tables': len(tables),
    }


# ---------------------------------------------------------------------------
# USD Generation
# ---------------------------------------------------------------------------

def sanitize_name(s: str) -> str:
    """Make a string safe for USD prim names."""
    s = re.sub(r'[^a-zA-Z0-9_]', '_', s)
    s = re.sub(r'_+', '_', s).strip('_')
    # Limit length to avoid huge prim names
    if len(s) > 60:
        s = s[:60]
    return s if s else 'unnamed'


def escape_usd_string(s: str) -> str:
    """Escape a string for USD string attribute value."""
    s = str(s).replace('\\', '\\\\').replace('"', '\\"')
    # Remove newlines and tabs
    s = s.replace('\n', ' ').replace('\r', ' ').replace('\t', ' ')
    # Limit length
    if len(s) > 200:
        s = s[:200] + '...'
    return s


def make_material(name, color, metallic=0.3, roughness=0.5):
    """Generate USD Material prim text."""
    r, g, b = color
    return f'''        def Material "{name}"
        {{
            token outputs:surface.connect = </AquilaSovereign/Materials/{name}/Shader.outputs:surface>
            def Shader "Shader"
            {{
                uniform token info:id = "UsdPreviewSurface"
                color3f inputs:diffuseColor = ({r:.3f}, {g:.3f}, {b:.3f})
                float inputs:metallic = {metallic:.2f}
                float inputs:roughness = {roughness:.2f}
                token outputs:surface
            }}
        }}'''


def make_box_mesh(name, dims, pos, mat_rel, bom_id='', subsystem='', weight=0, extra_attrs=None):
    """Generate USD Mesh box prim text."""
    w, h, d = dims
    x, y, z = pos
    lines = []
    lines.append(f'        def Mesh "{name}"')
    lines.append('        {')
    lines.append(f'            float3[] extent = [({-w/2:.1f}, {-h/2:.1f}, {-d/2:.1f}), ({w/2:.1f}, {h/2:.1f}, {d/2:.1f})]')
    lines.append('            int[] faceVertexCounts = [4, 4, 4, 4, 4, 4]')
    lines.append('            int[] faceVertexIndices = [0,1,2,3, 4,5,6,7, 0,4,5,1, 1,5,6,2, 2,6,7,3, 3,7,4,0]')
    lines.append(f'            rel material:binding = </AquilaSovereign/Materials/{mat_rel}>')
    lines.append(f'            point3f[] points = [({-w/2:.1f}, {-h/2:.1f}, {-d/2:.1f}), ({w/2:.1f}, {-h/2:.1f}, {-d/2:.1f}), ({w/2:.1f}, {h/2:.1f}, {-d/2:.1f}), ({-w/2:.1f}, {h/2:.1f}, {-d/2:.1f}), ({-w/2:.1f}, {-h/2:.1f}, {d/2:.1f}), ({w/2:.1f}, {-h/2:.1f}, {d/2:.1f}), ({w/2:.1f}, {h/2:.1f}, {d/2:.1f}), ({-w/2:.1f}, {h/2:.1f}, {d/2:.1f})]')
    if bom_id:
        lines.append(f'            string bomId = "{escape_usd_string(bom_id)}"')
    if subsystem:
        lines.append(f'            string subsystem = "{escape_usd_string(subsystem)}"')
    if weight:
        lines.append(f'            double weightGrams = {weight}')
    if extra_attrs:
        for k, v in extra_attrs.items():
            lines.append(f'            {k} = {v}')
    lines.append(f'            matrix4d xformOp:transform = ( (1, 0, 0, 0), (0, 1, 0, 0), (0, 0, 1, 0), ({x:.1f}, {y:.1f}, {z:.1f}, 1) )')
    lines.append('            token[] xformOpOrder = ["xformOp:transform"]')
    lines.append('        }')
    return '\n'.join(lines)


def make_cylinder_mesh(name, diam, height, pos, mat_rel, bom_id='', subsystem='', weight=0):
    """Generate USD Mesh cylinder prim text."""
    r = diam / 2
    n = 12  # segments
    pts = []
    # bottom ring
    for i in range(n):
        a = 2 * math.pi * i / n
        pts.append(f'({r * math.cos(a):.1f}, {r * math.sin(a):.1f}, {-height/2:.1f})')
    # top ring
    for i in range(n):
        a = 2 * math.pi * i / n
        pts.append(f'({r * math.cos(a):.1f}, {r * math.sin(a):.1f}, {height/2:.1f})')

    fvc = [4] * n + [n, n]  # side quads + bottom cap + top cap
    fvi = []
    for i in range(n):
        ni = (i + 1) % n
        fvi.extend([i, ni, ni + n, i + n])
    fvi.extend(list(range(n)))  # bottom
    fvi.extend(list(range(n, 2 * n)))  # top

    x, y, z = pos
    lines = []
    lines.append(f'        def Mesh "{name}"')
    lines.append('        {')
    lines.append(f'            float3[] extent = [({-r:.1f}, {-r:.1f}, {-height/2:.1f}), ({r:.1f}, {r:.1f}, {height/2:.1f})]')
    lines.append(f'            int[] faceVertexCounts = [{",".join(str(x) for x in fvc)}]')
    lines.append(f'            int[] faceVertexIndices = [{",".join(str(x) for x in fvi)}]')
    lines.append(f'            rel material:binding = </AquilaSovereign/Materials/{mat_rel}>')
    lines.append(f'            point3f[] points = [{", ".join(pts)}]')
    if bom_id:
        lines.append(f'            string bomId = "{escape_usd_string(bom_id)}"')
    if subsystem:
        lines.append(f'            string subsystem = "{escape_usd_string(subsystem)}"')
    if weight:
        lines.append(f'            double weightGrams = {weight}')
    lines.append(f'            matrix4d xformOp:transform = ( (1, 0, 0, 0), (0, 1, 0, 0), (0, 0, 1, 0), ({x:.1f}, {y:.1f}, {z:.1f}, 1) )')
    lines.append('            token[] xformOpOrder = ["xformOp:transform"]')
    lines.append('        }')
    return '\n'.join(lines)


def make_engineering_design_prim(name, spec, idx):
    """Generate a USD Scope prim for an engineering design spec from .docx."""
    lines = []
    lines.append(f'    def Scope "EngSpec_{idx:03d}_{sanitize_name(name)}"')
    lines.append('    {')
    lines.append(f'        string title = "{escape_usd_string(spec["title"][:200])}"')
    lines.append(f'        string source = "{escape_usd_string(spec["filename"])}"')
    if spec.get('doc_id'):
        lines.append(f'        string documentId = "{escape_usd_string(spec["doc_id"])}"')
    if spec.get('triz'):
        lines.append(f'        string trizPrinciple = "{escape_usd_string(spec["triz"])}"')
    if spec.get('date'):
        lines.append(f'        string designDate = "{escape_usd_string(spec["date"])}"')
    lines.append(f'        int paragraphCount = {spec["n_paragraphs"]}')
    lines.append(f'        int tableCount = {spec["n_tables"]}')

    # Add table data as string arrays
    for ti, table in enumerate(spec['tables'][:3]):  # max 3 tables
        header_cells = [str(c).replace('"', '').replace('\n', ' ')[:60] for c in table[0]]
        header_str = ", ".join(f'"{c}"' for c in header_cells)
        lines.append(f'        string[] table{ti}_header = [{header_str}]')
        for ri, row in enumerate(table[1:5]):  # max 5 rows
            cells = [str(c).replace('"', '').replace('\n', ' ')[:80] for c in row]
            row_str = ", ".join(f'"{c}"' for c in cells)
            lines.append(f'        string[] table{ti}_row{ri} = [{row_str}]')

    lines.append('    }')
    return '\n'.join(lines)


# ---------------------------------------------------------------------------
# Main USD Stage Builder
# ---------------------------------------------------------------------------

def build_complete_usd(components, specs, output_path, z_up=False):
    """Build the complete USD stage with BOM components + engineering specs."""
    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    up_axis = "Z" if z_up else "Y"

    lines = []
    lines.append('#usda 1.0')
    lines.append('(')
    lines.append(f'    defaultPrim = "AquilaSovereign"')
    lines.append(f'    doc = "AQUILA Sovereign Platform — Complete BOM + Engineering Specs USD. Generated {now}"')
    lines.append(f'    metersPerUnit = 0.001')
    lines.append(f'    upAxis = "{up_axis}"')
    lines.append(f'    framesPerSecond = 30')
    lines.append(f'    startTimeCode = 1')
    lines.append(f'    endTimeCode = 900')
    lines.append(')')
    lines.append('')
    lines.append('def Xform "AquilaSovereign"')
    lines.append('{')
    lines.append(f'    string mission = "AQUILA Sovereign — Orphan-Well Leak Detection + Bio-Intel"')
    lines.append(f'    string bomVersion = "Reconciled v2 7-5-26"')
    lines.append(f'    string platform = "Firefly Heavy-Lift Hexacopter"')
    lines.append(f'    int totalComponents = {len(components)}')
    lines.append(f'    int totalEngineeringSpecs = {len(specs)}')
    total_weight = sum(c.get('weight_g', 0) * c.get('qty', 1) for c in components)
    total_cost = sum(c.get('ext_cost', 0) for c in components)
    lines.append(f'    double totalWeightGrams = {total_weight:.0f}')
    lines.append(f'    double totalCostUSD = {total_cost:.0f}')
    lines.append('')

    # Metadata scope
    lines.append('    def Scope "Metadata"')
    lines.append('    {')
    lines.append(f'        string generatedBy = "Aquila Complete USD Converter v2.0"')
    lines.append(f'        string timestamp = "{now}"')
    lines.append(f'        string source = "Aquila_Master_BOM_Reconciled_v2 7-5-26.xlsx + {len(specs)} engineering spec .docx files"')
    lines.append(f'        string coordinateSystem = "Local frame (mm), {up_axis}-up"')
    lines.append(f'        string[] phases = ["Phase 1 — Operational", "Phase 2 — Firefly Hybrid Mount", "Phase 3 — Bio-Intel DNA Stack"]')
    lines.append('    }')
    lines.append('')

    # Materials scope
    lines.append('    def Scope "Materials"')
    lines.append('    {')

    # Generate unique materials per subsystem
    subsystem_colors = {
        'Power': (0.1, 0.6, 0.1),
        'Compute': (0.2, 0.2, 0.8),
        'Subsurface': (0.3, 0.8, 0.3),
        'Structure': (0.5, 0.5, 0.5),
        'EO/IR': (0.2, 0.2, 0.3),
        'Photogrammetry': (0.1, 0.1, 0.1),
        'Stereo Vision': (0.0, 0.6, 0.9),
        'NAV': (0.0, 0.8, 0.3),
        'Chemistry': (0.8, 0.2, 0.1),
        'Audio': (0.2, 0.0, 0.4),
        'Env': (0.0, 0.4, 0.6),
        'Thermal': (0.4, 0.4, 0.5),
        'Mount': (0.5, 0.5, 0.5),
        'Pod': (0.3, 0.3, 0.3),
        'Actuation': (0.8, 0.1, 0.1),
        'Control': (0.1, 0.1, 0.1),
        'Bio': (0.0, 0.7, 0.3),
        'Drone Platform': (0.15, 0.15, 0.18),
        'Gimbal': (0.6, 0.6, 0.65),
        'Payload Shell': (0.12, 0.12, 0.15),
        'Flight Battery': (0.1, 0.6, 0.1),
        'Payload Battery': (0.1, 0.5, 0.2),
        'Compute Core': (0.2, 0.2, 0.8),
        'Carrier Board': (0.15, 0.15, 0.6),
        'Subsurface Radar': (0.3, 0.8, 0.3),
    }

    materials_made = set()
    for comp in components:
        subsys = comp.get('Subsystem', 'Default')
        if subsys not in materials_made:
            color = subsystem_colors.get(subsys, (0.5, 0.5, 0.5))
            mat_name = f"{sanitize_name(subsys)}_Mat"
            lines.append(make_material(mat_name, color))
            materials_made.add(subsys)

    lines.append('    }')
    lines.append('')

    # BOM Components scope
    lines.append('    def Scope "BOM_Components"')
    lines.append('    {')

    # Auto-position components in a grid if no position data
    grid_x = 0
    grid_y = 0
    grid_spacing = 400

    for comp in components:
        name = sanitize_name(comp.get('ID', 'comp'))
        comp_name = comp.get('Component / Model', name)[:60]
        dims = comp.get('dims_mm', (50, 50, 50))
        weight = comp.get('weight_g', 0)
        subsys = comp.get('Subsystem', 'Default')
        qty = comp.get('qty', 1)

        mat_name = f"{sanitize_name(subsys)}_Mat"

        # Position components in a grid
        pos = (grid_x, grid_y, 0)

        # Determine shape
        dims_str = comp.get('Dimensions (mm)', '')
        is_cylinder = dims_str.startswith('Ø') or dims_str.startswith('o')

        for q in range(qty):
            qname = f"{name}_{q:02d}" if qty > 1 else name
            qpos = (pos[0] + q * (dims[0] + 50), pos[1], pos[2])

            if is_cylinder:
                lines.append(make_cylinder_mesh(
                    qname, dims[0], dims[2], qpos, mat_name,
                    bom_id=comp.get('ID', ''),
                    subsystem=subsys,
                    weight=weight,
                ))
            else:
                lines.append(make_box_mesh(
                    qname, dims, qpos, mat_name,
                    bom_id=comp.get('ID', ''),
                    subsystem=subsys,
                    weight=weight,
                    extra_attrs={
                        'string componentName': f'"{escape_usd_string(comp_name)}"',
                    },
                ))

        # Advance grid position
        grid_x += grid_spacing
        if grid_x > 2000:
            grid_x = 0
            grid_y += grid_spacing

    lines.append('    }')
    lines.append('')

    # Engineering Specs scope
    lines.append('    def Scope "EngineeringSpecs"')
    lines.append('    {')
    lines.append(f'        string description = "Engineering design innovations from {len(specs)} .docx spec files"')
    lines.append('')

    for idx, spec in enumerate(specs):
        if 'error' in spec:
            lines.append(f'    # Skipped {spec["filename"]}: {spec["error"]}')
            continue
        title_short = spec['title'][:60] if spec['title'] else spec['filename'][:60]
        lines.append(make_engineering_design_prim(title_short, spec, idx))
        lines.append('')

    lines.append('    }')
    lines.append('')

    lines.append('}')
    lines.append('')

    # Write output
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))

    return len(components), len(specs)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Aquila Complete USD Converter")
    parser.add_argument('--output', default=str(OUTPUT_DEFAULT), help='Output .usda path')
    parser.add_argument('--z-up', action='store_true', help='Use Z-up (for Isaac Sim)')
    parser.add_argument('--base-dir', default=str(BASE_DIR), help='Source directory')
    args = parser.parse_args()

    base = Path(args.base_dir)
    output = Path(args.output)

    print(f"[Aquila USD] Source: {base}")
    print(f"[Aquila USD] Output: {output}")

    # 1. Parse BOM xlsx
    print(f"\n[1/3] Parsing BOM spreadsheet...")
    bom_file = base / "Aquila_Master_BOM_Reconciled_v2 7-5-26 (version 1).xlsx"
    if not bom_file.exists():
        # Find any xlsx
        xlsx_files = list(base.glob("*.xlsx"))
        if xlsx_files:
            bom_file = xlsx_files[0]
        else:
            print("ERROR: No BOM xlsx found", file=sys.stderr)
            sys.exit(1)

    components = parse_bom_xlsx(bom_file)
    print(f"  → {len(components)} BOM components parsed")
    for c in components[:5]:
        print(f"    {c.get('ID', '?'):10s} {c.get('Subsystem', '?'):20s} {c.get('Component / Model', '?')[:40]}")

    # 2. Parse all .docx engineering specs
    print(f"\n[2/3] Parsing engineering spec .docx files...")
    docx_files = sorted(base.glob("*.docx"))
    # Skip temp files
    docx_files = [f for f in docx_files if not f.name.startswith('~$')]
    # Also check subdirectories
    sub_docx = sorted((base / "Autonomous Scientific Sampling Subsystem").glob("*.docx"))
    docx_files.extend(sub_docx)

    specs = []
    for f in docx_files:
        spec = parse_docx_spec(f)
        specs.append(spec)
        status = "OK" if 'error' not in spec else f"ERROR: {spec['error']}"
        print(f"  {f.name[:55]:55s} → {status} ({spec.get('n_paragraphs', 0)} para, {spec.get('n_tables', 0)} tbl)")

    print(f"  → {len(specs)} engineering specs parsed ({len([s for s in specs if 'error' not in s])} OK)")

    # 3. Build complete USD stage
    print(f"\n[3/3] Building complete USD stage...")
    n_comp, n_specs = build_complete_usd(components, specs, output, z_up=args.z_up)

    file_size = output.stat().st_size
    print(f"\n✓ Complete USD stage generated: {output}")
    print(f"  Components: {n_comp}")
    print(f"  Engineering specs: {n_specs}")
    print(f"  File size: {file_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
